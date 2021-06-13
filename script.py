# importing libraries
from docx import Document
import requests
from colorama import init
from termcolor import colored
init()
import time
import os
import firebase_admin
from firebase_admin import credentials
from firebase_admin import firestore



cred = credentials.Certificate("./serviceAccountKey.json")
firebase_admin.initialize_app(cred)

db = firestore.client()

docs = db.collection(u'data').stream()








# creating a .docx file
doc = Document()

os.system('cls||clear')
print(colored('\n--------------------\ automate MS word with python /--------------------', 'red'))
print(colored('----------------------------\ by omar chatin /----------------------------\n', 'red'))

count = 0

reqst = input('type "r" to get random data from api or type "s" to get real data : ')




# input how many students do you want to get








if reqst=='s':
    for dc in docs:
        students_number = 0

    table = doc.add_table(rows=students_number+1,cols=5)
    for doc in docs:
        person = doc.to_dict()
        count +=1
        dt_table = table.rows[count].cells
        dt_table[0].text = person["name"]
        dt_table[1].text = person["gender"]
        dt_table[2].text = person["email"]
        dt_table[3].text = person["phone"]
        dt_table[4].text = person["State"]
        print(person["name"])
        print(count)

else:
    print(colored('error occured \n write "r" to get random data or type "s" to get real data \n\n','white', 'on_red'))









# styling the table
table.style = 'Table Grid'

th_table = table.rows[0].cells
th_table[0].text = 'name'
th_table[1].text = 'gender'
th_table[2].text = 'email'
th_table[3].text = 'phone'
th_table[4].text = 'location'

print(colored('collecting data from api ...','green'))

# requesting data
data = requests.get(f'https://randomuser.me/api/?results={students_number}')


data = data.json()['results']



if reqst=='r':
    table = doc.add_table(rows=students_number+1,cols=5)

    students_number = int(input(colored('how many people do you want to get ? : ','green')))
    # creating the table
    table = doc.add_table(rows=students_number+1,cols=5)

    for i in data:
        count +=1
        person={
            "name" : i["name"]["first"] + i["name"]["last"],
            "gender" : i["gender"],
            "email" : i["email"],
            "phone" : i["phone"],
            "location" : i["location"]["country"] + ' / ' +  i["location"]["city"],
        }

        print(colored(f'prson {count} data: ', 'yellow'))
        print(person, '\n')



        dt_table = table.rows[count].cells
        dt_table[0].text = person["name"]
        dt_table[1].text = person["gender"]
        dt_table[2].text = person["email"]
        dt_table[3].text = person["phone"]
        dt_table[4].text = person["location"]



# changing thew page layout
from docx.shared import Inches, Cm
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

# saving the file
doc.save('demo.docx')

time.sleep(2)
os.system('cls')

print(colored('---- finished ----', 'red'))
time.sleep(2)

print(colored('\n now openning the word file ...', 'red'))
time.sleep(2)

# openning the .docx file
os.system('start demo.docx')