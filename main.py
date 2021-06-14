# importing libraries/////////////////////                
                                    #/////
from docx import Document           #/////                    
import requests                     #/////        
from colorama import init           #/////                    
from termcolor import colored       #/////                        
init()                              #/////
import time                         #/////    
import os                           #/////    
import firebase_admin               #/////                
from firebase_admin import credentials#///                         
from firebase_admin import firestore#/////                                
from docx.shared import Inches, Cm, Pt
import matplotlib.pyplot as plt
import numpy as np





gender_plt = ['male', 'female']
gender_sizes = [0,0]

age_plt = ['less than 19','more than 19']
age_sizes = [0,0]




#createing the table and word file////////
#/////////////////////////////////////////
                                    #/////
word = Document()                   #///// 
table = word.add_table(rows=1,cols=6)#////
table.style = 'Medium Shading 1 Accent 6'#
th_cells = table.rows[0].cells      #/////
th_cells[0].text = 'name'           #/////
th_cells[1].text = 'gender'         #/////
th_cells[2].text = 'email'          #/////
th_cells[3].text = 'phone'          #/////
th_cells[4].text = 'State'          #/////
th_cells[5].text = 'Age'          #/////
                                    #/////
# page layout/////////////////////////////
#///////////////////////////////////////// 
                                    #/////
sections = word.sections            #/////
for section in sections:            #/////
    section.top_margin = Cm(0.5)    #/////
    section.bottom_margin = Cm(0.5) #/////
    section.left_margin = Cm(0.5)   #/////
    section.right_margin = Cm(0.5)  #/////
                                    #/////
#/////////////////////////////////////////
#/////////////////////////////////////////
count = 0
err=False 

os.system('cls||clear')
print(colored('\n---------------------\ automate MS word with python /--------------------', 'red'))
print(colored('----------------------------\ by omar chatin /----------------------------\n\n', 'red'))



rqst = input(colored('type "rdm" to get random data from api or type "db" to get real data from database : ','green'))



# getting data from firebase
if rqst=='db':
    print(colored('fetching data from database...\n', 'red'))
    time.sleep(4)
    cred = credentials.Certificate("./serviceAccountKey.json")
    firebase_admin.initialize_app(cred)
    db = firestore.client()
    docs = db.collection(u'data').stream()

    print(docs)
    for prs in docs:
        person = prs.to_dict()
        a = table.add_row().cells
        a[0].text = person['name']
        a[1].text = person['gender']
        a[2].text = person['email']
        a[3].text = person['phone']
        a[4].text = person['State']
        a[5].text = person['age']


        count+=1
        print(colored(f'person -{count}- : \n','yellow'))
        print(person)
        if person['gender']=='male':
            gender_sizes[0]+=1
        elif person['gender']=='female':
            gender_sizes[1]+=1
        
        if int(person['age'])<19:
            age_sizes[0]+=1
        elif int(person['age'])>=19:
            age_sizes[1]+=1
        else:
            err=False







# getting data from random user api
elif rqst=='rdm':

    try:
        num = int(input(colored('how many people do you want to get ? : ','green')))
    except:
        print(colored('error occured , make sure you filled the inputs correctly .','red'))
        time.sleep(3)
        exit()

    print(colored('\nfetching data from randomuser API...', 'red'))
    time.sleep(4)
    data = requests.get(f'https://randomuser.me/api/?results={num}')
    data = data.json()['results']
    for i in data:
        person={
            "name" : i["name"]["first"] + i["name"]["last"],
            "gender" : i["gender"],
            "email" : i["email"],
            "phone" : i["phone"],
            "location" : i["location"]["country"] + ' / ' +  i["location"]["city"],
            "age" : i['dob']['age'],
        }

        a = table.add_row().cells
        a[0].text = person['name']
        a[1].text = person['gender']
        a[2].text = person['email']
        a[3].text = person['phone']
        a[4].text = person['location']
        a[5].text = str(person['age'])

        count+=1
        print(colored(f'person -{count}- : \n','yellow'))
        print(person)
        if person['gender']=='male':
            gender_sizes[0]+=1
        elif person['gender']=='female':
            gender_sizes[1]+=1


else:
    time.sleep(3)
    print(colored('\nerror occured , make sure you filled the inputs correctly .','red'))
    exit()





time.sleep(3)
os.system('cls')

print(colored('\n---- finished processing ----', 'red'))
time.sleep(1)



print(colored('\n drawing charts with matplotlib...', 'red'))
time.sleep(2)





age_sizes[0] = age_sizes[0] * 100/count
age_sizes[1] = age_sizes[1] * 100/count

fig1, ax1= plt.subplots()
ax1.pie(gender_sizes,labels=gender_plt)
plt.savefig('./data/gender.png')
pic1 = word.add_picture('./data/gender.png',width=Inches(4.0))

if not(rqst=='rdm'):
    fig2, ax2= plt.subplots()
    ax2.pie(age_sizes,labels=age_plt)
    plt.savefig('./data/age.png')
    pic2 = word.add_picture('./data/age.png',width=Inches(4.0))








# plt.show()
print(colored('\n now openning the word file automatically...', 'red'))
time.sleep(2)
word.save('./data/demo2.docx')
os.system('start data/demo2.docx')